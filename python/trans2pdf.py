import os
from django.conf import settings
from docx import Document
from docx.shared import Inches
from docx.oxml.ns import qn
import fitz

from concurrent import futures
import logging

import grpc
import helloworld_pb2
import helloworld_pb2_grpc


class Greeter(helloworld_pb2_grpc.GreeterServicer):

    def SayHello(self, request, context):
        c = open(file="/media/bzl/DATA3/BZLproject/00_AGV/forcompare/mywebserver/example/ml.pdf", mode="rb+")
        # file_context = request.name
        doc = fitz.open(stream=c) # 打开一个pdf文件
        rotate = int(0) # 设置图片的旋转角度        
        trans = fitz.Matrix(2.0, 2.0).prerotate(theta=0)
        # file_context = 
        # chuli
        return helloworld_pb2.HelloReply(message='Hello, %s!' % request.name)


def serve():
    port = '50051'
    server = grpc.server(futures.ThreadPoolExecutor(max_workers=10))
    helloworld_pb2_grpc.add_GreeterServicer_to_server(Greeter(), server)
    server.add_insecure_port('[::]:' + port)
    server.start()
    print("Server started, listening on " + port)
    server.wait_for_termination()

import requests   # pip intasll requests
import random
import hashlib
import re
import django
import time
import urllib.request
import urllib.parse
import json

# os.environ.setdefault("DJANGO_SETTINGS_MODULE", "trans2pdf.settings")
# django.setup()
def baidu_translate(contest):    
    '''实现有道翻译的接口'''
    url = 'http://fanyi.youdao.com/translate?smartresult=dict&smartresult=rule&sessionFrom=https://www.baidu.com/link'
    data = {
        'from':'AUTO',
        'to':'AUTO',
        'smartresult':'dict',
        'client':'fanyideskweb',
        'salt':'1500092479607',
        'sign':'d9f9a3aa0a7b34241b3fe30505e5d436',
        'doctype':'json',
        'version':'2.1',
        'keyfrom':'fanyi.web',
        'action':'FY_BY_CL1CKBUTTON',
        'typoResult':'true'}
    data['i'] = content.replace('\n','')
    data = urllib.parse.urlencode(data).encode('utf-8')
    wy = urllib.request.urlopen(url,data)
    html = wy.read().decode('utf-8')
    ta = json.loads(html)
    res = ta['translateResult'][0][0]['tgt']
    return res

# 百度翻译方法
def baidu_translate2(content):
    print(content)
    print("test==")
    if len(content) > 4891:
        return '输入请不要超过4891个字符！'
    print("test==2")
    salt = str(random.randint(0, 50))
    # 申请网站 http://api.fanyi.baidu.com/api/trans
    appid = '20221026001417491' # 这里写你自己申请的
    secretKey = '6lCldAVUsF09IRI6KDan'# 这里写你自己申请的
    sign = appid + content + salt + secretKey
    sign = hashlib.md5(sign.encode(encoding='UTF-8')).hexdigest()
    head = {'q': f'{content}',
            'from': 'en',
            'to': 'zh',
            'appid': f'{appid }',
            'salt': f'{salt}',
            'sign': f'{sign}'}
    print("test==3")
    j = requests.get('http://api.fanyi.baidu.com/api/trans/vip/translate', head)
    print("test==4")
    print(j.json())
    res = j.json()['trans_result'][0]['dst']
    print("test==5")
    res = re.compile('[\\x00-\\x08\\x0b-\\x0c\\x0e-\\x1f]').sub(' ', res)
    print(res)
    return res

# 正则匹配参考文献
def is_reference(target):
    return re.match(r'references', target, re.I)

# 正则匹配图片标注
def is_figure(target):
    return re.match(r'fig\..\.', target, re.I)

if __name__ == '__main__':
    c = open(file="/media/bzl/DATA3/BZLproject/00_AGV/forcompare/mywebserver/example/ml.pdf", mode="rb+")
    # file_context = request.name
    print(type(c))
    file_name = "test.pdf"
    t0 = time.time()
    cur_pdf = fitz.open(c) # 待翻译的pdf
    new_pdf = fitz.open() # 翻译完成后要写入的pdf
    new_docx = Document() # 翻译完成后要写入的docx
    new_docx.styles["Normal"].font.name = "宋体"
    new_docx.styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    
    i = 0 # 定义页面数的递增
    bytes_array = 0

    try:
        for cur_page in cur_pdf:
            img_list = cur_page.get_images()
            img_count = 0
            for img in img_list:
                pix_temp1 = fitz.Pixmap(cur_pdf, img[0])
                if img[1]:
                    pix_temp2 = fitz.Pixmap(cur_pdf, img[1])
                    pix_temp = fitz.Pixmap(pix_temp1)
                    pix_temp.setAlpha(pix_temp2.samples)
                else:
                    pix_temp = pix_temp1
                
                img_count += 1
                new_name = "图片{}.png".format(img_count) # 生成图片名称
                pix_temp.save(new_name)
                pix_temp = None
            print("当前正在翻译第{}页...".format(int(str(cur_page).split(" ")[1]) + 1))
            blocks = cur_page.get_text("blocks",clip=None, flags=None, textpage=None, sort=False)
            new_page = new_pdf.new_page(pno=-1, width=cur_page.mediabox_size[0], height=cur_page.mediabox_size[1])
            img = new_page.new_shape()
            disp = fitz.Rect(cur_page.cropbox_position, cur_page.cropbox_position)
            croprect = cur_page.rect + disp
            begin = (0, 0, 0, 0)
            end = (0, 0, 0, 0)
            flag = 0
            reference_flag = 0
            blocks.append((1, 2, 3, 6))
            content = ""
            img_count = 0
            fonts = 9
            print("===len:", len(blocks))
            for num in range(len(blocks)):
                # 如果是本页面最后一个块,直接结束,因为最后一个是方便计算自己添加的。
                if num == len(blocks) - 1:
                    break
                # 如果这个块里放的是图像.
                print("===test1")
                if blocks[num][-1] == 1:
                    print('图像:::',blocks[num][4])
                    img_count += 1
                    # 图片要放置位置的坐标
                    img_r = blocks[num][:4]
                    # 当前页面第几个图片的位置
                    image_path = os.path.join('图片{}.png'.format(img_count))
                    # 输入流
                    img = open(image_path, "rb").read()
                    # 输入到新的pdf页面对应位置
                    new_page.insert_image(img_r, stream=img, keep_proportion=True)
                    # 设置图片保存的宽度
                    new_docx.add_picture(image_path, width=Inches(3))
                    # 输入到新的pdf之后就移除
                    os.remove(image_path)
                    continue

                print("===test2")
                # 设置默认字体大小以及位置
                if i == 0:  # 当前是第一页的话
                    # 一般论文前面的标题,作者,机构名等要居中
                    print("===test3")    
                    if num == 0 or num == 1:
                        fonts = 15
                        text_pos = fitz.TEXT_ALIGN_CENTER
                    elif num == 2:
                        fonts = 10
                        text_pos = fitz.TEXT_ALIGN_CENTER
                    elif num == 3:
                        fonts = 10
                        text_pos = fitz.TEXT_ALIGN_CENTER
                    # 设置文字在当前矩阵中的位置靠左排列
                    else:
                        fonts = 10
                        text_pos = fitz.TEXT_ALIGN_LEFT
                # 设置文字在当前矩阵中的位置靠左排列
                else:
                    print("===test4")
                    fonts = 10
                    text_pos = fitz.TEXT_ALIGN_LEFT
                # 目的为了记录起始块坐标
                if num == 0:
                    print("===test5")    
                    begin = blocks[0][:4]
                    content = blocks[0][4].replace("\n", " ")
                # 矩形块，b[0]b[1]为左上角的坐标，b[2]b[3]为右下角的坐标
                print("===test6")
                r = fitz.Rect(blocks[num][:4])
                # 如果不是倒数第一个块，则进入此循环
                print("===test7")    
                if num < len(blocks) - 1:
                    print("===test8")
                    # 两个块y轴距离很近的话，这里以1.0为界，这里判断当前数的右下角的坐标y值
                    if (abs(blocks[num + 1][1] - blocks[num][3]) <= 1.0 and abs(
                            blocks[num + 1][1] - blocks[num][3]) >= 0):
                        # 当前块在参考文献之后
                        print("===test8.1")
                        if reference_flag == 1:
                            print("===test9")
                            trans_pragraph = blocks[num][4].replace("\n", " ")
                            res = baidu_translate(trans_pragraph).replace(' ', '')
                            new_page.insert_textbox(r, res, fontname="song", fontfile=os.path.join('SimSun.ttf'),
                                                   fontsize=7, align=text_pos)  #
                        # 其它情况
                        else:
                            print("===test10")
                            flag = 1  #
                            # 记录最后的矩形坐标，目的为了取出最后的右下角坐标点
                            end = blocks[num + 1][:4]
                            content += blocks[num + 1][4].replace("\n", " ")
                            # print('content::',content)

                    # 两个块y轴距离远的的时候
                    else:
                        if flag == 1:
                            print("===test11")
                            # img.drawRect(fitz.Rect(end[0],begin[1],end[2],end[3]))
                            res = baidu_translate(content).replace(' ', '')  # 翻译结果去掉汉字中的空格
                            new_docx.add_paragraph(res)  # 添加到新的docx文档中
                            # print('content:',content)
                            # print(res)
                            # fitz.Rect(end[0],begin[1],end[2],end[3])为新扩展的矩形框坐标
                            if begin[2] > end[2]:  # 如果起始点的右下角x坐标小于结束点的右下角x坐标
                                new_page.insert_textbox(fitz.Rect(end[0], begin[1], begin[2], end[3]), res, fontname="song",
                                                    fontfile=os.path.join('SimSun.ttf'),
                                                    fontsize=fonts, align=text_pos)
                            else:
                                new_page.insert_textbox(fitz.Rect(end[0], begin[1], end[2], end[3]), res, fontname="song",
                                                    fontfile=os.path.join('SimSun.ttf'),
                                                    fontsize=fonts, align=text_pos)
                            flag = 0
                        else:
                            print("===test12")
                            # img.drawRect(r)
                            trans_pragraph = blocks[num][4].replace("\n", " ")  # 将待翻译的句子换行换成空格
                            if is_figure(trans_pragraph.replace(' ','')):  # 将该块的判断是否是图片标注
                                res = baidu_translate(trans_pragraph).replace(' ', '')  # 翻译结果去掉汉字中的空格
                                print("res", res)
                                new_page.insert_textbox(r, res, fontname="song", fontfile=os.path.join('SimSun.ttf'),
                                                    fontsize=7, align=fitz.TEXT_ALIGN_CENTER)
                            # 标记在这里之后的都是参考文献
                            elif is_reference(trans_pragraph.replace(' ','')):
                                print("===test13")
                                reference_flag = 1
                                new_page.insert_textbox(r, '参考文献', fontname="song", fontfile=os.path.join('SimSun.ttf'),
                                                    fontsize=fonts, align=text_pos)
                            else:
                                print("===test14=====")
                                # 翻译结果去掉汉字中的空格
                                res = baidu_translate(trans_pragraph).replace(' ', '')
                                print("res:", res)
                                # 添加到新的docx文档中
                                new_docx.add_paragraph(res)
                                if reference_flag == 1:
                                    print("===test15")
                                    new_page.insert_textbox(r, res, fontname="song", fontfile=os.path.join('SimSun.ttf'),
                                                           fontsize=7, align=text_pos)  #
                                else:
                                    print("===test16")
                                    new_page.insert_textbox(r, res, fontname="song", fontfile=os.path.join('SimSun.ttf'),
                                                    fontsize=fonts, align=text_pos)  #
                        # 记录起始矩形坐标
                        begin = blocks[num + 1][:4]
                        try:
                            content = blocks[num + 1][4].replace("\n", " ")
                            # print('content:::',content)
                        except:
                            pass
                            #print('记录content失败！')
                        # img.finish(width=0.3)
                        # img.commit()
            i += 1
    except Exception as error:
        print("翻译过程异常:", error)

    # 文件保存
    # 翻译后的pdf保存路径
    new_pdf_name = os.path.join(file_name)
    # 翻译后的docx保存路径
    new_docx_name = os.path.join(file_name[:-4] + '.docx')
    new_docx.save(new_docx_name)
    new_pdf.save(new_pdf_name, garbage=4, deflate=True, clean=True)
    t1 = time.time()
    # trans = fitz.Matrix(2.0, 2.0).prerotate(theta=0)    
    # for index in range(doc.page_count):
    #     page = doc[index]
    #     pm = page.get_pixmap(matrix=trans, dpi=None, colorspace=fitz.csRGB, clip=None, alpha=False, annots=True)
    #     new_full_name = name.split(".")[0]
    #     if doc.page_count > 1 :
    #         pm.save(filename="%s%s.pdf" % (new_full_name, index), output="png")
    #     else:
    #         pm.save(filename="%s.pdf" % (new_full_name), output="png")
    
    print("pdf convert done")
